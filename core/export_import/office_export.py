    
import os
import io

import uuid
import pandas as pd

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt

from openpyxl import load_workbook 
from openpyxl.drawing.image import Image

from .html_to_docx import add_html
from . import Downloader
    
    

class OfficeDownloader():
    
    @staticmethod
    def __add_fig_to_docx(document, fig, width_percent):
        with io.BytesIO() as io_fig:
            fig.write_image(file=io_fig, format='png', engine="kaleido")
            section = document.sections[0]
            width = width_percent*section.page_width
            document.add_picture(io_fig, width=width)
        return document
    
    @staticmethod
    def __insert_table(document, tablename, comment_under_table ,df):
        
        document.add_heading(tablename, level=1)
        t = document.add_table(df.shape[0]+1, df.shape[1])
            
        for j in range(df.shape[1]):
            res = f"{df.columns[j]}"
            t.cell(0,j).text = res

        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                if type(df.iloc[i,j]) == str:
                    res = df.iloc[i,j]
                else:
                    res = f"{df.iloc[i,j]:_.0f}".replace("_", ".")
                
                t.cell(i+1,j).text = res
            
        t.style = 'Table Grid'
        document.add_paragraph(comment_under_table)
        return document

    @staticmethod
    def __modify_styles(document, lable: str ,styles: dict):
        styles = document.styles[lable]
        color = styles.get('color')
        styles.font.color.rgb = RGBColor.from_string(color)
        
        size = styles.get('size')
        if size:
            styles.font.size = Pt(size)

        name = styles.get('name')
        if name:
            styles.font.name = name

        attributes = styles.get('attribures')
        if attributes:
            if 'italic' in attributes:
                styles.font.italic = True
            else:
                styles.font.italic = False
            if 'bold' in attributes:
                styles.font.bold = True
            else:
                styles.font.bold = False
            if 'underline' in attributes:
                styles.font.underline = True
        return document
    
    
    
    @staticmethod
    def __insert_markdown(document, markdown):
        html_tree = markdown.markdown(markdown)
        paragraph = document.add_paragraph()
        add_html(paragraph, html_tree)

    @staticmethod
    def create_docx_for_download(input_list, styles_list):
        document = Document()
   
        for item in styles_list:
            lable = styles_list.get('lable')
            if lable:
                document = OfficeDownloader.__modify_styles(document, lable=lable, styles=item)
            if lable == 'Title':
                bottom_color = item.get('bottom_color')
                if bottom_color:
                    title_styles = document.styles['Title']
                    bottom = title_styles.element.xpath("./w:pPr/w:pBdr/w:bottom")[0]
                    bottom.set(qn("w:color"), bottom_color)
                    bottom.attrib.pop(qn("w:themeColor"))
        
        for input_item in input_list:
            input_type = input_item.get('type')
            if input_type == 'text':
                style = input_item.get('style')
                text = input_item.get('text')
                document.add_paragragh(text, style=style)
            if input_type == 'fig':
                fig = input_item.get('fig')
                width_percent = input_item.get('width')
                document = OfficeDownloader.__add_fig_to_docx(document, fig, width_percent)
            if input_type == 'table':
                tablename = input_item.get('table_headline')
                comment_under_table = input_item.get('comment')
                df = input_item.get('df')
                document = OfficeDownloader.__insert_table(document, tablename, comment_under_table ,df)
            if input_type == 'markdown':     
                markdown = input_item.get('markdown')
                document = OfficeDownloader.__insert_markdown(document, markdown)
        
        with io.BytesIO() as io_docx:
            document.save(io_docx)
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"     
            download = Downloader.__create_downloader_href(io_docx, mimetype)  
        return download



    @staticmethod
    def create_xlsx_for_download(list_of_df: list):
        '''
        list of arguments in the format [{'df':Dataframe, 'sheetname' : sheetname}, ]
        produces downloadable xlsx file
        '''
        
        xlsx_io = io.BytesIO()
        with pd.ExcelWriter(xlsx_io, engine="xlsxwriter") as writer:
            for item in  list_of_df:
                df = item.get('df')
                df.to_excel(writer, sheet_name=item.get('sheetname'))    

                
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        download = Downloader._create_downloader_href(xlsx_io, mimetype=mimetype)
        
        return download 


    # needs approval to implement it fully
    # @staticmethod
    # def create_ppt_for_download(data, data_corrected, chapter_num, fig):
    #     document = Downloader.__create_docx_with_table(data, data_corrected, chapter_num)

        
    #     with io.BytesIO() as io_docx:
    #         document.save(io_docx)
    #         mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"     
    #         download = Downloader.__create_downloader_href(io_docx, mimetype)  
    #     return download